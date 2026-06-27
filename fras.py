"""Streamlit entrypoint for FRAS — File & Report Automation System."""

from __future__ import annotations

import io
import json
import os
import time
import uuid
from pathlib import Path
from typing import Any

import streamlit as st
from dotenv import load_dotenv

import hashlib

from db import (
    add_comment,
    delete_file,
    get_all_files,
    get_comments,
    get_dashboard_stats,
    get_extracted_data,
    get_file,
    get_file_by_hash,
    get_files_by_ids,
    get_connection,
    init_db,
    insert_extracted_data,
    insert_file,
    search_files,
    update_extracted_structured_data,
    update_file_category,
    update_file_error,
    update_file_flagged,
    update_file_reviewed,
    update_file_status,
)
from openrouter_client import (
    ask_documents,
    detect_risks,
    extract_document_content,
    generate_insights,
    generate_report_content,
    semantic_search,
)
from enhanced_ask_client import ask_documents_reasoned

load_dotenv()

# ---------------------------------------------------------------------------
# Config
# ---------------------------------------------------------------------------
STORAGE_DIR = Path("storage/files")
STORAGE_DIR.mkdir(parents=True, exist_ok=True)

ALLOWED_EXTENSIONS = {".pdf", ".docx", ".txt", ".png", ".jpg", ".jpeg"}
MIME_MAP = {
    ".pdf": "application/pdf",
    ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ".txt": "text/plain",
    ".png": "image/png",
    ".jpg": "image/jpeg",
    ".jpeg": "image/jpeg",
}

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _compute_file_hash(uploaded_file: Any) -> str:
    """Compute SHA-256 hash of uploaded file bytes."""
    hasher = hashlib.sha256()
    for chunk in iter(lambda: uploaded_file.read(8192), b""):
        hasher.update(chunk)
    uploaded_file.seek(0)
    return hasher.hexdigest()


def _save_uploaded_file(uploaded_file: Any) -> tuple[str, str]:
    """Save uploaded file to storage and return (saved_path, file_type)."""
    ext = Path(uploaded_file.name).suffix.lower()
    if ext not in ALLOWED_EXTENSIONS:
        raise ValueError(f"Unsupported file type: {ext}")

    unique_name = f"{uuid.uuid4().hex}_{uploaded_file.name}"
    saved_path = STORAGE_DIR / unique_name

    with open(saved_path, "wb") as f:
        f.write(uploaded_file.getbuffer())

    return str(saved_path), ext


def _process_single_file(file_id: int, file_path: str, file_type: str) -> None:
    """Run OpenRouter extraction for one file and store results."""
    try:
        update_file_status(file_id, "processing")
        result = extract_document_content(file_path, file_type)

        risks = result.get("risks", json.dumps([]))
        structured_data = result.get("structured_data", json.dumps({}))

        insert_extracted_data(
            file_id=file_id,
            summary=result.get("summary"),
            key_points=result.get("key_points"),
            document_type=result.get("document_type"),
            entities=result.get("entities"),
            raw_response=result.get("raw_response"),
            risks=risks,
            structured_data=structured_data,
        )
        # Store category from Gemini output
        doc_type = result.get("document_type")
        if doc_type:
            update_file_category(file_id, doc_type)

        # Feature 6: Workflow Automation — apply rules
        _apply_workflow_rules(file_id)

        update_file_status(file_id, "done")
    except Exception as exc:
        update_file_status(file_id, "failed")
        update_file_error(file_id, str(exc))
        if "processing_errors" not in st.session_state:
            st.session_state.processing_errors = []
        st.session_state.processing_errors.append(
            f"Failed to process {Path(file_path).name}: {exc}"
        )


def _apply_workflow_rules(file_id: int) -> None:
    """
    Feature 6: Apply simple rule-based automation after document processing.
    Rules:
      - If document_type == "invoice" → category = "Finance"
      - If risks not empty → flagged = True
    """
    file = get_file(file_id)
    if not file:
        return
    extracted = get_extracted_data(file_id)
    if not extracted:
        return

    # Rule 1: Invoices go to Finance
    doc_type = extracted.get("document_type", "")
    if doc_type and doc_type.lower() == "invoice":
        update_file_category(file_id, "Finance")

    # Rule 2: Flag if risks exist
    risks_raw = extracted.get("risks", "[]")
    try:
        risks_list = json.loads(risks_raw) if risks_raw else []
    except (json.JSONDecodeError, TypeError):
        risks_list = []
    if risks_list:
        update_file_flagged(file_id, True)


def _get_clean_text(file_id: int) -> str:
    """Get document text for risk re-detection."""
    file = get_file(file_id)
    if not file:
        return ""
    from openrouter_client import _extract_text_from_file
    return _extract_text_from_file(file["filepath"], file["file_type"])


def _parse_json_field(value: Any, default: Any = None) -> Any:
    """Safely parse a JSON field from the database."""
    if value is None:
        return default if default is not None else []
    if isinstance(value, (list, dict)):
        return value
    try:
        return json.loads(value)
    except (json.JSONDecodeError, TypeError):
        return default if default is not None else []


# ---------------------------------------------------------------------------
# UI — Tabs
# ---------------------------------------------------------------------------

def render_upload_tab() -> None:
    st.header("Upload Documents")
    st.caption("Supported formats: PDF, DOCX, TXT, PNG, JPG")

    if st.session_state.get("processing_errors"):
        for err in st.session_state.processing_errors:
            st.error(err)
        if st.button("Dismiss errors"):
            st.session_state.processing_errors = []
            st.rerun()

    uploaded_files = st.file_uploader(
        "Choose files",
        type=["pdf", "docx", "txt", "png", "jpg", "jpeg"],
        accept_multiple_files=True,
    )

    if not uploaded_files:
        return

    if st.button("Upload & Process", type="primary"):
        progress_bar = st.progress(0, text="Starting upload...")
        total = len(uploaded_files)

        for idx, uploaded_file in enumerate(uploaded_files, start=1):
            progress_bar.progress(
                (idx - 1) / total,
                text=f"Saving file {idx} of {total}: {uploaded_file.name}",
            )
            try:
                file_hash = _compute_file_hash(uploaded_file)
                existing = get_file_by_hash(file_hash)
                if existing:
                    st.warning(
                        f"This file was already uploaded as {existing['filename']} on {existing['upload_date'][:10]}. Skipping."
                    )
                    continue

                saved_path, file_type = _save_uploaded_file(uploaded_file)
                file_id = insert_file(
                    filename=uploaded_file.name,
                    filepath=saved_path,
                    file_type=file_type,
                )
                with get_connection() as conn:
                    conn.execute(
                        "UPDATE files SET file_hash = ? WHERE id = ?",
                        (file_hash, file_id),
                    )
            except Exception as exc:
                st.error(f"Error saving {uploaded_file.name}: {exc}")
                continue

            progress_bar.progress(
                (idx - 0.5) / total,
                text=f"Processing file {idx} of {total}: {uploaded_file.name}",
            )
            _process_single_file(file_id, saved_path, file_type)

        progress_bar.progress(1.0, text="All files processed!")
        time.sleep(0.5)
        progress_bar.empty()
        st.success("Upload and processing complete!")
        st.rerun()


def render_library_tab(role: str = "Owner") -> None:
    st.header("Document Library")

    all_files = get_all_files()

    # Filter: Show flagged only
    show_flagged_only = st.checkbox("Show flagged only", value=False)
    if show_flagged_only:
        all_files = [f for f in all_files if f.get("flagged")]

    # Category filter
    categories = sorted({f.get("category") or "Uncategorized" for f in all_files if f.get("category")})
    category_options = ["All"] + categories
    selected_category = st.selectbox("Filter by category", options=category_options)

    # Feature 7: Replace keyword search with intent-based search
    search_query = st.text_input("Search documents (natural language)", placeholder="e.g. 'Show me all invoices with amounts over 1000'")

    if search_query.strip():
        try:
            # Gather all done documents with extracted data
            done_files = [f for f in all_files if f["status"] == "done"]
            doc_list = []
            for f in done_files:
                ext = get_extracted_data(f["id"])
                if ext and ext.get("summary"):
                    doc_list.append({
                        "filename": f["filename"],
                        "summary": ext.get("summary", ""),
                        "key_points": ext.get("key_points", "[]"),
                    })
            if doc_list:
                matches = semantic_search(search_query.strip(), doc_list)
                if matches:
                    files = [f for f in all_files if f["filename"] in matches]
                else:
                    files = []
            else:
                files = all_files
        except Exception as exc:
            st.warning(f"Intent-based search failed, falling back to keyword search: {exc}")
            files = search_files(search_query.strip())
    else:
        files = all_files

    if selected_category != "All":
        files = [f for f in files if (f.get("category") or "Uncategorized") == selected_category]

    if not files:
        st.info("No documents found. Upload some files first!")
        return

    display_rows = []
    for file in files:
        extracted = get_extracted_data(file["id"])
        summary = extracted["summary"][:120] + "..." if extracted and extracted.get("summary") else "—"
        flagged_icon = "🚩 " if file.get("flagged") else ""
        display_rows.append(
            {
                "ID": file["id"],
                "Filename": f"{flagged_icon}{file['filename']}",
                "Uploaded": file["upload_date"][:10],
                "Type": file["file_type"],
                "Status": file["status"],
                "Category": file.get("category") or "Uncategorized",
                "Summary": summary,
            }
        )

    st.dataframe(display_rows, use_container_width=True, hide_index=True)

    # Expandable detail view
    st.subheader("Document Details")
    selected_id = st.number_input("Enter document ID to view details", min_value=1, step=1)

    col_view, col_retry, col_delete = st.columns(3)
    with col_view:
        view_clicked = st.button("View Details", key="view_details_btn")
    with col_retry:
        retry_clicked = st.button("Retry", key=f"retry_{selected_id}", disabled=(role != "Owner"))
    with col_delete:
        delete_clicked = st.button("Delete", key=f"delete_{selected_id}", type="secondary", disabled=(role != "Owner"))

    if delete_clicked and role == "Owner":
        st.session_state[f"confirm_delete_{selected_id}"] = True
        st.rerun()

    if st.session_state.get(f"confirm_delete_{selected_id}") and role == "Owner":
        st.warning(f"Are you sure you want to delete document ID {selected_id}?")
        col_confirm, col_cancel = st.columns(2)
        with col_confirm:
            if st.button("Yes, delete", key=f"confirm_delete_yes_{selected_id}", type="primary"):
                try:
                    delete_file(selected_id)
                    st.success("Document deleted.")
                    st.session_state[f"confirm_delete_{selected_id}"] = False
                    st.rerun()
                except Exception as exc:
                    st.error(f"Delete failed: {exc}")
        with col_cancel:
            if st.button("Cancel", key=f"confirm_delete_no_{selected_id}"):
                st.session_state[f"confirm_delete_{selected_id}"] = False
                st.rerun()

    if retry_clicked and role == "Owner":
        file = get_file(selected_id)
        if not file:
            st.error("Document not found.")
            return
        with st.spinner(f"Retrying processing for {file['filename']}..."):
            _process_single_file(selected_id, file["filepath"], file["file_type"])
        st.success("Retry complete. Check status above.")
        st.rerun()

    if view_clicked:
        file = get_file(selected_id)
        if not file:
            st.error("Document not found.")
            return

        extracted = get_extracted_data(selected_id)

        st.write(f"**Filename:** {file['filename']}" + (" 🚩" if file.get("flagged") else ""))
        st.write(f"**Uploaded:** {file['upload_date']}")
        st.write(f"**Type:** {file['file_type']}")
        st.write(f"**Status:** {file['status']}")
        st.write(f"**Category:** {file.get('category') or 'Uncategorized'}")

        if file.get("last_error"):
            st.error(f"**Last Error:** {file['last_error']}")

        if extracted:
            st.write("---")
            st.write(f"**Document Type:** {extracted.get('document_type', 'N/A')}")
            st.write(f"**Summary:** {extracted.get('summary', 'N/A')}")

            # Key Points
            key_points = extracted.get("key_points")
            if key_points:
                try:
                    kp_list = json.loads(key_points)
                    st.write("**Key Points:**")
                    for kp in kp_list:
                        st.write(f"- {kp}")
                except json.JSONDecodeError:
                    st.write(f"**Key Points:** {key_points}")

            # Entities
            entities = extracted.get("entities")
            if entities:
                try:
                    ent_list = json.loads(entities)
                    st.write("**Entities:**")
                    st.write(", ".join(ent_list))
                except json.JSONDecodeError:
                    st.write(f"**Entities:** {entities}")

            # Feature 3: Risks display
            risks = extracted.get("risks")
            if risks:
                try:
                    risks_list = json.loads(risks) if risks else []
                except (json.JSONDecodeError, TypeError):
                    risks_list = []
                if risks_list:
                    st.warning("**Risks Detected ⚠️**")
                    for risk in risks_list:
                        st.write(f"- ⚠️ {risk}")
                else:
                    st.success("No risks detected for this document.")
            else:
                st.info("Risk analysis not run for this document.")

            # Feature 5: Structured Data display
            structured_data = extracted.get("structured_data")
            if structured_data:
                try:
                    sd = json.loads(structured_data) if structured_data else {}
                except (json.JSONDecodeError, TypeError):
                    sd = {}
                if sd and isinstance(sd, dict) and any(v for v in sd.values()):
                    st.write("**Structured Data:**")
                    for key, value in sd.items():
                        if value:
                            display_key = key.replace("_", " ").title()
                            if isinstance(value, list):
                                st.write(f"- **{display_key}:** {', '.join(str(v) for v in value)}")
                            else:
                                st.write(f"- **{display_key}:** {value}")

            if role == "Owner":
                with open(file["filepath"], "rb") as f:
                    st.download_button(
                        label="Download Original File",
                        data=f.read(),
                        file_name=file["filename"],
                        mime="application/octet-stream",
                    )

            # Feature 8: Collaboration - Comments
            st.write("---")
            st.subheader("Collaboration")

            # Mark as reviewed
            col_review, col_status = st.columns(2)
            with col_review:
                reviewed = file.get("reviewed", False)
                label = "Mark as Reviewed" if not reviewed else "Reviewed ✅"
                if st.button(label, key=f"review_{selected_id}"):
                    update_file_reviewed(selected_id, not reviewed)
                    st.rerun()
            with col_status:
                if file.get("reviewed"):
                    st.success("This document has been reviewed.")
                else:
                    st.info("This document has not been reviewed yet.")

            # Comments section
            st.write("**Comments**")
            comment_text = st.text_input("Add a comment", key=f"comment_input_{selected_id}")
            if st.button("Add Comment", key=f"add_comment_{selected_id}"):
                if comment_text.strip():
                    add_comment(selected_id, comment_text.strip())
                    st.success("Comment added!")
                    st.rerun()
                else:
                    st.warning("Please enter a comment.")

            comments = get_comments(selected_id)
            if comments:
                for c in comments:
                    ts = c.get("timestamp", "")[:19].replace("T", " ")
                    st.caption(f"{ts}")
                    st.write(c["comment"])
                    st.divider()
            else:
                st.write("No comments yet.")
        else:
            st.warning("No extracted data available for this document.")


def render_reports_tab() -> None:
    st.header("Generate Report")

    files = get_all_files()
    done_files = [f for f in files if f["status"] == "done"]

    if not done_files:
        st.warning("No processed documents available. Upload and process files first.")
        return

    options = {f"{f['id']} — {f['filename']}": f["id"] for f in done_files}
    selected_labels = st.multiselect(
        "Select documents to include in the report",
        options=list(options.keys()),
    )

    if not selected_labels:
        st.info("Select at least one document to generate a report.")
        return

    selected_ids = [options[label] for label in selected_labels]

    if st.button("Generate Report", type="primary"):
        with st.spinner("Generating report with Gemini..."):
            summaries = []
            for file_id in selected_ids:
                file = get_file(file_id)
                extracted = get_extracted_data(file_id)
                if file and extracted:
                    summaries.append(
                        {
                            "filename": file["filename"],
                            "summary": extracted.get("summary", ""),
                            "key_points": extracted.get("key_points", ""),
                        }
                    )

            if not summaries:
                st.error("No extracted data found for selected documents.")
                return

            try:
                report_md = generate_report_content(summaries)
            except Exception as exc:
                st.error(f"Report generation failed: {exc}")
                return

        st.subheader("Generated Report")
        st.markdown(report_md)

        st.download_button(
            label="Download as Markdown (.md)",
            data=report_md,
            file_name="fras_report.md",
            mime="text/markdown",
        )

        try:
            from docx import Document

            doc = Document()
            doc.add_heading("FRAS Generated Report", level=0)
            for line in report_md.splitlines():
                if line.startswith("# "):
                    doc.add_heading(line[2:], level=1)
                elif line.startswith("## "):
                    doc.add_heading(line[3:], level=2)
                elif line.startswith("- "):
                    doc.add_paragraph(line[2:], style="List Bullet")
                elif line.strip():
                    doc.add_paragraph(line)

            docx_bytes = io.BytesIO()
            doc.save(docx_bytes)
            docx_bytes.seek(0)

            st.download_button(
                label="Download as Word (.docx)",
                data=docx_bytes,
                file_name="fras_report.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        except Exception as exc:
            st.warning(f"DOCX export unavailable: {exc}")


# ---------------------------------------------------------------------------
# Feature 1: Ask Across Documents
# ---------------------------------------------------------------------------

def render_ask_tab() -> None:
    st.header("Ask Across Documents")
    st.caption("Ask natural language questions about your documents. Enable reasoning mode for feasibility analysis, step-by-step reasoning, and evidence-backed answers using full document content.")

    files = get_all_files()
    done_files = [f for f in files if f["status"] == "done"]

    if not done_files:
        st.warning("No processed documents available. Upload and process files first.")
        return

    # Multi-select documents
    options = {f"{f['id']} — {f['filename']}": f["id"] for f in done_files}
    selected_labels = st.multiselect(
        "Select documents to query",
        options=list(options.keys()),
        key="ask_doc_select",
    )

    if not selected_labels:
        st.info("Select at least one document to ask questions about.")
        return

    selected_ids = [options[label] for label in selected_labels]

    # Reasoning mode toggle
    enable_reasoning = st.checkbox(
        "Enable reasoning mode",
        value=False,
        key="ask_reasoning_toggle",
        help="When enabled, includes full document text for step-by-step reasoning and feasibility analysis (e.g., 'Given the weather data, is it possible to schedule an outdoor event?'). Uses more context tokens.",
    )

    question = st.text_input(
        "Your question",
        placeholder="e.g., What are the total amounts mentioned in these documents?",
        key="ask_question",
    )

    if not question.strip():
        return

    if st.button("Ask", type="primary", key="ask_button"):
        with st.spinner("Analyzing documents to answer your question..."):
            from openrouter_client import _extract_text_from_file

            # Gather document context
            doc_contexts = []
            for file_id in selected_ids:
                file = get_file(file_id)
                extracted = get_extracted_data(file_id)
                if file and extracted:
                    doc_item = {
                        "filename": file["filename"],
                        "summary": extracted.get("summary", ""),
                        "key_points": extracted.get("key_points", "[]"),
                        "entities": extracted.get("entities", "[]"),
                        "structured_data": extracted.get("structured_data", "{}"),
                        "risks": extracted.get("risks", "[]"),
                    }
                    # If reasoning mode, include full document text
                    if enable_reasoning:
                        doc_item["full_text"] = _extract_text_from_file(
                            file["filepath"], file["file_type"]
                        )
                    doc_contexts.append(doc_item)

            if not doc_contexts:
                st.error("No extracted data found for selected documents.")
                return

            try:
                if enable_reasoning:
                    result = ask_documents_reasoned(doc_contexts, question.strip())
                else:
                    result = ask_documents(doc_contexts, question.strip())
            except Exception as exc:
                st.error(f"Failed to process question: {exc}")
                return

        # Display answer
        st.subheader("Answer")
        st.info(result.get("answer", "Could not generate an answer."))

        # Display reasoning (only in reasoning mode)
        reasoning = result.get("reasoning")
        if reasoning:
            st.subheader("Reasoning")
            st.markdown(reasoning)

        # Display feasibility (only in reasoning mode)
        feasibility = result.get("feasibility")
        if feasibility and feasibility != "not_applicable":
            feasibility_map = {
                "possible": ("✅ Possible", "green"),
                "not_possible": ("❌ Not Possible", "red"),
                "insufficient_data": ("⚠️ Insufficient Data", "orange"),
            }
            label, color = feasibility_map.get(
                feasibility, (f"⚪ {feasibility.replace('_', ' ').title()}", "gray")
            )
            st.markdown(f"**Feasibility:** <span style='color:{color};font-weight:bold'>{label}</span>", unsafe_allow_html=True)

        # Display sources
        sources = result.get("sources", [])
        if sources:
            st.subheader("Sources")
            for i, source in enumerate(sources, 1):
                doc_name = source.get("document", "Unknown")
                evidence = source.get("evidence", "No evidence provided.")
                with st.expander(f"Source {i}: {doc_name}", expanded=(i <= 3)):
                    st.write(evidence)


# ---------------------------------------------------------------------------
# Feature 2: Insight Engine
# ---------------------------------------------------------------------------

def render_insights_tab() -> None:
    st.header("Insight Engine")
    st.caption("Analyze multiple documents together to extract cross-document insights, recurring issues, trends, and risks.")

    files = get_all_files()
    done_files = [f for f in files if f["status"] == "done"]

    if not done_files:
        st.warning("No processed documents available. Upload and process files first.")
        return

    options = {f"{f['id']} — {f['filename']}": f["id"] for f in done_files}
    selected_labels = st.multiselect(
        "Select documents to analyze together",
        options=list(options.keys()),
        key="insight_doc_select",
    )

    if not selected_labels:
        st.info("Select at least one document to generate insights.")
        return

    selected_ids = [options[label] for label in selected_labels]

    if st.button("Generate Insights", type="primary", key="insight_button"):
        with st.spinner("Generating cross-document insights... (this may take a moment)"):
            doc_contexts = []
            for file_id in selected_ids:
                file = get_file(file_id)
                extracted = get_extracted_data(file_id)
                if file and extracted:
                    entities_raw = extracted.get("entities", "[]")
                    doc_contexts.append({
                        "filename": file["filename"],
                        "summary": extracted.get("summary", ""),
                        "key_points": extracted.get("key_points", "[]"),
                        "entities": entities_raw,
                    })

            if not doc_contexts:
                st.error("No extracted data found for selected documents.")
                return

            try:
                insights = generate_insights(doc_contexts)
            except Exception as exc:
                st.error(f"Failed to generate insights: {exc}")
                return

        # Display Issues
        st.subheader("Top Recurring Issues")
        issues = insights.get("issues", [])
        if issues:
            for issue in issues:
                st.write(f"- {issue}")
        else:
            st.write("No recurring issues identified.")

        # Display Entities grouped
        st.subheader("Key Entities Grouped")
        entities = insights.get("entities", {})
        if entities and isinstance(entities, dict):
            for category, items in entities.items():
                if items:
                    st.write(f"**{category.title()}:** {', '.join(items)}")
        else:
            st.write("No entities identified.")

        # Display Trends
        st.subheader("Trends & Patterns")
        trends = insights.get("trends", [])
        if trends:
            for trend in trends:
                st.write(f"- {trend}")
        else:
            st.write("No trends identified.")

        # Display Risks
        st.subheader("Observed Risks")
        risk_items = insights.get("risks", [])
        if risk_items:
            for risk in risk_items:
                st.warning(f"⚠️ {risk}")
        else:
            st.success("No risks identified.")


# ---------------------------------------------------------------------------
# Feature 4: Smart Dashboard
# ---------------------------------------------------------------------------

@st.cache_data(ttl=60)
def _get_cached_dashboard_stats() -> dict:
    """Cached dashboard stats to keep UI fast."""
    return get_dashboard_stats()


def render_dashboard_tab() -> None:
    st.header("Smart Dashboard")
    st.caption("Overview of your document collection.")

    stats = _get_cached_dashboard_stats()

    # Metrics at top
    col1, col2, col3 = st.columns(3)
    with col1:
        st.metric("Total Documents", stats["total"])
    with col2:
        st.metric("Documents with Risks ⚠️", stats["with_risks"])
    with col3:
        st.metric("Flagged Documents 🚩", stats["flagged"])

    # Bar chart for document types
    st.subheader("Documents by Type")
    by_type = stats.get("by_type", [])
    if by_type:
        chart_data = {"Type": [], "Count": []}
        for item in by_type:
            chart_data["Type"].append(item["type"])
            chart_data["Count"].append(item["count"])
        st.bar_chart(chart_data, x="Type", y="Count", use_container_width=True)
    else:
        st.info("No documents uploaded yet.")

    # Line chart for uploads over time
    st.subheader("Uploads Over Time")
    uploads = stats.get("uploads_over_time", [])
    if uploads:
        chart_data = {"Date": [], "Count": []}
        for item in uploads:
            chart_data["Date"].append(item["date"])
            chart_data["Count"].append(item["count"])
        st.line_chart(chart_data, x="Date", y="Count", use_container_width=True)
    else:
        st.info("No upload activity yet.")


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main() -> None:
    st.set_page_config(page_title="FRAS", page_icon="📄", layout="wide")
    init_db()

    # Clear cache on refresh
    st.cache_data.clear()

    # Role selector (MVP simplification — no real auth)
    with st.sidebar:
        st.title("FRAS")
        st.caption("File & Report Automation System")
        role = st.selectbox("Role", ["Owner", "Viewer"], index=0)
        st.caption("Viewer mode hides upload/delete/retry actions.")
        st.divider()
        st.caption("v2.0 · All 8 features")

    # Create tabs
    tab_upload, tab_library, tab_reports, tab_ask, tab_insights, tab_dashboard = st.tabs(
        ["Upload", "Library", "Reports", "Ask", "Insights", "Dashboard"]
    )

    with tab_upload:
        if role == "Owner":
            render_upload_tab()
        else:
            st.info("Viewer role: upload is disabled. Switch to Owner in the sidebar to upload files.")

    with tab_library:
        render_library_tab(role=role)

    with tab_reports:
        render_reports_tab()

    with tab_ask:
        render_ask_tab()

    with tab_insights:
        render_insights_tab()

    with tab_dashboard:
        render_dashboard_tab()


if __name__ == "__main__":
    main()